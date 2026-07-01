import docx

def build_opsi1():
    doc = docx.Document("templates/jurnal_komputer_sinta.docx")
    
    # 0: Judul
    doc.paragraphs[0].text = "{{ judul_artikel }}"
    
    # 2: Penulis
    doc.paragraphs[2].text = "{{ nama_penulis }}"
    
    # 3: Afiliasi
    doc.paragraphs[3].text = "{{ afiliasi }}"
    
    # Delete 4 to 8
    for i in range(8, 3, -1):
        try:
            p = doc.paragraphs[i]
            p.text = ""
            p._element.getparent().remove(p._element)
        except IndexError:
            pass
            
    # Now indices shift by 5 (8-4+1 = 5)
    # The original 9 (Corresponding author) is now 9-5 = 4
    # Wait, 8 was deleted, so index 4 is the new [9]
    # Let's just find the text "Corresponding author"
    
    corr_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "Corresponding author" in p.text:
            corr_idx = i
            break
            
    if corr_idx != -1:
        doc.paragraphs[corr_idx].text = "Email korespondensi: {{ email_korespondensi }}"
        # Delete the next line which was the email placeholder
        p_to_del = doc.paragraphs[corr_idx + 1]
        p_to_del.text = ""
        p_to_del._element.getparent().remove(p_to_del._element)

    # Abstrak
    abs_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "Abstrak" in p.text and len(p.text) < 20:
            abs_idx = i
            break
            
    if abs_idx != -1:
        doc.paragraphs[abs_idx].text = "ABSTRAK"
        # delete next 2 lines
        for i in range(abs_idx + 2, abs_idx, -1):
            p_to_del = doc.paragraphs[i]
            p_to_del.text = ""
            p_to_del._element.getparent().remove(p_to_del._element)
            
        doc.paragraphs[abs_idx + 1].text = "{{ abstrak }}"
        
        # Kata kunci is the line after next
        # Let's find "Kata Kunci:"
        key_idx = -1
        for i, p in enumerate(doc.paragraphs):
            if "Kata Kunci:" in p.text:
                key_idx = i
                break
        
        if key_idx != -1:
            doc.paragraphs[key_idx].text = "Kata kunci: {{ kata_kunci }}"
            
    # Find PENDAHULUAN
    pend_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "PENDAHULUAN" in p.text:
            pend_idx = i
            break
            
    if pend_idx != -1:
        doc.paragraphs[pend_idx].text = "{% for bab in daftar_bab %}"
        doc.paragraphs[pend_idx+1].text = "{{ bab.judul_bab }}"
        doc.paragraphs[pend_idx+2].text = "{{ bab.isi_bab }}"
        doc.paragraphs[pend_idx+3].text = "{% endfor %}"
        
    # Delete everything between bab loop and references
    ref_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "DAFTAR PUSTAKA" in p.text:
            ref_idx = i
            break
            
    if ref_idx != -1 and pend_idx != -1:
        for i in range(ref_idx - 1, pend_idx + 3, -1):
            try:
                p = doc.paragraphs[i]
                p.text = ""
                p._element.getparent().remove(p._element)
            except IndexError:
                pass
                
    # Now references is shifted
    ref_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "DAFTAR PUSTAKA" in p.text:
            ref_idx = i
            break
            
    if ref_idx != -1:
        doc.paragraphs[ref_idx].text = "DAFTAR PUSTAKA"
        doc.paragraphs[ref_idx+1].text = "{% for ref in daftar_referensi %}"
        doc.paragraphs[ref_idx+2].text = "{{ ref.teks_sitasi }}"
        doc.paragraphs[ref_idx+3].text = "{% endfor %}"
        
        for i in range(len(doc.paragraphs) - 1, ref_idx + 3, -1):
            try:
                p = doc.paragraphs[i]
                p.text = ""
                p._element.getparent().remove(p._element)
            except IndexError:
                pass

    doc.save("templates/jurnal_komputer_sinta_v2.docx")
    print("Saved templates/jurnal_komputer_sinta_v2.docx")

if __name__ == "__main__":
    build_opsi1()
