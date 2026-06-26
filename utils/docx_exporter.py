import re
from pathlib import Path
import docx

from utils.token_counter import truncate_to_tokens

def extract_template_text(template_path: str) -> str:
    if not template_path:
        return ""
    p_path = Path(template_path)
    text = ""
    if p_path.suffix.lower() == ".docx":
        try:
            doc = docx.Document(template_path)
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(lines)
        except Exception:
            text = ""
    else:
        try:
            text = p_path.read_text(encoding="utf-8")
        except Exception:
            text = ""
            
    # Truncate to a safe token limit to prevent exceeding API limits
    return truncate_to_tokens(text, 1000)


def export_markdown_to_docx(md_path: str, docx_path: str, template_path: str = None) -> None:
    md_content = Path(md_path).read_text(encoding="utf-8")
    
    # 1. Parse metadata (title, abstract, keywords) from markdown frontmatter
    title = ""
    keywords_list = []
    
    fm_match = re.match(r"^---\s*\n(.*?)\n---\s*\n", md_content, re.DOTALL)
    if fm_match:
        fm_text = fm_match.group(1)
        t_match = re.search(r"^title:\s*[\"']?(.*?)[\"']?\s*$", fm_text, re.MULTILINE)
        if t_match:
            title = t_match.group(1)
        k_match = re.search(r"^keywords:\s*\[?(.*?)\]?\s*$", fm_text, re.MULTILINE)
        if k_match:
            k_raw = k_match.group(1)
            keywords_list = [k.strip(" '\"") for k in k_raw.split(",")]
            
    # Parse abstract from body
    abstract_match = re.search(r"## Abstrak\s*\n(.*?)(?=\n##|$)", md_content, re.DOTALL)
    abstract_text = abstract_match.group(1).strip() if abstract_match else ""
    
    keywords_str = ", ".join(keywords_list)
    
    doc = None
    keyword_idx = -1
    
    # 2. If a docx template is provided, load it and perform smart inline replacement
    if template_path and Path(template_path).suffix.lower() == ".docx":
        try:
            doc = docx.Document(template_path)
            
            # Helper to replace text of a paragraph while keeping its runs/formatting intact
            def replace_paragraph_text_preserving_format(p, new_text):
                if not p.runs:
                    p.add_run(new_text)
                else:
                    p.runs[0].text = new_text
                    # Remove all extra runs so we don't have overlapping old text
                    for r in list(p.runs)[1:]:
                        p._element.remove(r._element)
            
            # Find title paragraph
            # The title is usually the first non-empty paragraph that is not a header/journal metadata
            title_p = None
            for p in doc.paragraphs:
                cleaned = p.text.strip()
                if cleaned and len(cleaned) > 5:
                    if not any(x in cleaned.lower() for x in ["jurnal", "issn", "volume", "halaman", "vol.", "no."]):
                        title_p = p
                        break
            if title_p and title:
                replace_paragraph_text_preserving_format(title_p, title)
                
            # Find and replace Abstrak / Abstract and Keywords / Kata Kunci
            for idx, p in enumerate(list(doc.paragraphs)):
                txt = p.text.lower()
                if "kata kunci" in txt or "keywords" in txt:
                    keyword_idx = idx
                    if "kata kunci" in txt:
                        replace_paragraph_text_preserving_format(p, f"Kata Kunci: {keywords_str}")
                    else:
                        replace_paragraph_text_preserving_format(p, f"Keywords: {keywords_str}")
                elif "abstrak" in txt or "abstract" in txt:
                    # If the paragraph itself contains the abstract text (longer paragraph)
                    if len(p.text.strip()) > 50:
                        prefix = "Abstrak—" if "abstrak" in txt else "Abstract—"
                        replace_paragraph_text_preserving_format(p, f"{prefix}{abstract_text}")
                    # If it's a short header, check the next paragraph for the body
                    elif idx + 1 < len(doc.paragraphs):
                        next_p = doc.paragraphs[idx + 1]
                        if len(next_p.text.strip()) > 50:
                            prefix = "Abstrak—" if "abstrak" in txt else "Abstract—"
                            replace_paragraph_text_preserving_format(next_p, f"{prefix}{abstract_text}")
                            
            # Delete all paragraphs after the keyword paragraph to clean the template body
            # but keep the section break containing the multi-column formatting
            if keyword_idx != -1:
                for p in list(doc.paragraphs)[keyword_idx + 1:]:
                    p._element.getparent().remove(p._element)
            else:
                # If keywords section not found, fallback to clearing all paragraphs
                for p in list(doc.paragraphs):
                    p._element.getparent().remove(p._element)
        except Exception as e:
            print(f"[Warning] Failed to use docx template: {e}")
            doc = docx.Document()
    else:
        doc = docx.Document()
        
    # 3. If no template or template failed, apply standard styling
    if not template_path or not doc.paragraphs:
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = docx.shared.Pt(11)
        
    # 4. Parse the markdown sections and append to the docx document
    in_yaml = False
    for line in md_content.splitlines():
        trimmed = line.strip()
        
        # Skip YAML frontmatter
        if trimmed == "---":
            in_yaml = not in_yaml
            continue
        if in_yaml:
            continue
            
        # Empty lines
        if not trimmed:
            continue
            
        # Skip Abstrak section since it was already replaced in the template header
        # (Only if we are using a template and successfully replaced the header elements)
        if template_path and keyword_idx != -1:
            # Skip Title heading and Abstrak section in the body list
            if trimmed.startswith("# "):
                continue
            if trimmed.startswith("## Abstrak") or trimmed.startswith("## Abstract"):
                # We skip the heading and the subsequent abstract paragraph
                continue
            # Check if this paragraph is the abstract paragraph that follows the heading
            # (We skip it by checking if it's the exact abstract text)
            if abstract_text and trimmed.startswith(abstract_text[:min(50, len(abstract_text))]):
                continue
                
        # Blockquotes (warnings)
        if trimmed.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = docx.shared.Inches(0.4)
            run = p.add_run(trimmed.lstrip(">").strip())
            run.italic = True
            continue
            
        # Headings
        if trimmed.startswith("# "):
            try:
                doc.add_heading(trimmed[2:], level=1)
            except Exception:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = docx.shared.Pt(12)
                p.paragraph_format.space_after = docx.shared.Pt(6)
                run = p.add_run(trimmed[2:])
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = docx.shared.Pt(16)
            continue
        elif trimmed.startswith("## "):
            try:
                doc.add_heading(trimmed[3:], level=2)
            except Exception:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = docx.shared.Pt(12)
                p.paragraph_format.space_after = docx.shared.Pt(4)
                run = p.add_run(trimmed[3:])
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = docx.shared.Pt(13)
            continue
        elif trimmed.startswith("### "):
            try:
                doc.add_heading(trimmed[4:], level=3)
            except Exception:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = docx.shared.Pt(6)
                p.paragraph_format.space_after = docx.shared.Pt(2)
                run = p.add_run(trimmed[4:])
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = docx.shared.Pt(11)
            continue
            
        # Bullet list items
        if trimmed.startswith("- ") or trimmed.startswith("* "):
            try:
                p = doc.add_paragraph(style='List Bullet')
                _add_formatted_text(p, trimmed[2:])
            except Exception:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = docx.shared.Inches(0.25)
                p.paragraph_format.space_after = docx.shared.Pt(3)
                p.add_run("• ")
                _add_formatted_text(p, trimmed[2:])
            continue
            
        # Standard paragraphs
        p = doc.add_paragraph()
        p.paragraph_format.space_after = docx.shared.Pt(6)
        _add_formatted_text(p, line)
        
    doc.save(docx_path)


def _add_formatted_text(paragraph, text: str):
    # Split by bold and italic syntax: **bold** or *italic*
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)
