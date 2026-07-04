"""
Template DOCX Injector
-----------------------
Prinsip: Deterministic Layout Engine

Tugasnya sempit dan jelas:
1. Terima structured_doc dict (dari doc_extractor agent atau dari pipeline AI)
2. Buka template .docx
3. Inject konten ke placeholder template menggunakan python-docx
4. Simpan file hasil

Aturan desain:
- File ini tidak pernah memanggil LLM
- File ini tidak menghasilkan konten baru
- Semua keputusan visual (font, margin, spasi, kolom) dikontrol template .docx
- Jika tag placeholder tidak ditemukan, backend melaporkan error yang ramah (tidak crash)
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Optional
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn


# ── Placeholder tag patterns yang didukung ──────────────────────────────────
# Template .docx boleh pakai {{TITLE}}, {{ABSTRACT}}, {{KEYWORDS}},
# serta {{SECTIONS}} sebagai blok tempat bab-bab disuntikkan.
# Jika memakai docxtpl Jinja2, gunakan flow yang lama (export_to_docx).
SUPPORTED_TAGS = ["{{TITLE}}", "{{ABSTRACT}}", "{{KEYWORDS}}", "{{SECTIONS}}"]


def inject_into_template(
    structured_doc: dict,
    template_path: str,
    output_path: str,
) -> tuple[str, list[str]]:
    """
    Inject konten dari structured_doc ke dalam template .docx.
    
    Args:
        structured_doc: dict dengan keys title, abstract, keywords, sections
                        (output dari doc_extractor atau pipeline AI)
        template_path:  path ke file .docx template jurnal
        output_path:    path untuk menyimpan file hasil

    Returns:
        tuple(output_path, warnings)  
        warnings = list of tag issues yang ditemukan (misal: tag tidak ada di template)
    """
    warnings: list[str] = []

    if not template_path or not Path(template_path).exists():
        # Tidak ada template → buat dokumen kosong dengan format standar
        warnings.append("Tidak ada template yang diunggah. Menggunakan format dokumen standar.")
        _export_plain_docx(structured_doc, output_path)
        return output_path, warnings

    # Validasi template dan baca teks
    try:
        doc = Document(template_path)
        template_raw_text = "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        warnings.append(f"File template rusak atau format tidak valid. Menggunakan format dokumen standar.")
        _export_plain_docx(structured_doc, output_path)
        return output_path, warnings

    if "{{" in template_raw_text and "{%" in template_raw_text:
        # Template dengan Jinja2 → pakai docxtpl (cara lama yang sudah terbukti)
        _export_via_jinja(structured_doc, template_path, output_path)
        return output_path, warnings

    # Template dengan placeholder sederhana {{TAG}} → deterministic injection
    found_tags = [tag for tag in SUPPORTED_TAGS if tag in template_raw_text]
    missing_tags = [tag for tag in SUPPORTED_TAGS if tag not in template_raw_text]

    if missing_tags:
        for tag in missing_tags:
            warnings.append(
                f"Tag {tag} tidak ditemukan di template. "
                f"Pastikan template memiliki placeholder {tag} agar konten dapat disuntikkan."
            )

    if not found_tags:
        # Tidak ada tag sama sekali → fallback ke mode smart replacement
        warnings.append(
            "Template tidak memiliki placeholder tag. "
            "Menggunakan mode penggantian cerdas (smart replacement) berdasarkan deteksi konten."
        )
        _export_smart_replace(structured_doc, template_path, output_path)
        return output_path, warnings

    # Ada placeholder tag → lakukan injection deterministik
    _export_via_placeholder(structured_doc, template_path, output_path, warnings)
    return output_path, warnings


def _read_docx_text(path: str) -> str:
    """Baca seluruh teks dari .docx untuk cek keberadaan tag."""
    try:
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception:
        return ""


def _export_via_placeholder(
    structured_doc: dict,
    template_path: str,
    output_path: str,
    warnings: list[str],
) -> None:
    """
    Ganti placeholder {{TAG}} di template dengan konten dari structured_doc.
    Mempertahankan formatting runs di paragraf template.
    """
    doc = Document(template_path)

    title = structured_doc.get("title", "")
    abstract = structured_doc.get("abstract", "")
    keywords = structured_doc.get("keywords", [])
    sections = structured_doc.get("sections", [])

    keywords_str = "; ".join(keywords) if keywords else ""

    for para in doc.paragraphs:
        full_text = para.text

        if "{{TITLE}}" in full_text:
            _replace_placeholder_in_para(para, "{{TITLE}}", title)

        if "{{ABSTRACT}}" in full_text:
            _replace_placeholder_in_para(para, "{{ABSTRACT}}", abstract)

        if "{{KEYWORDS}}" in full_text:
            _replace_placeholder_in_para(para, "{{KEYWORDS}}", keywords_str)

    # Untuk {{SECTIONS}}, kita perlu menemukan paragraf placeholder dan menggantinya
    # dengan beberapa paragraf baru (heading + konten tiap bab)
    sections_para_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "{{SECTIONS}}" in para.text:
            sections_para_idx = i
            break

    if sections_para_idx is not None:
        # Hapus paragraf placeholder
        placeholder_para = doc.paragraphs[sections_para_idx]
        placeholder_parent = placeholder_para._element.getparent()
        placeholder_index = list(placeholder_parent).index(placeholder_para._element)
        placeholder_parent.remove(placeholder_para._element)

        # Insert paragraf bab satu per satu setelah posisi placeholder
        inserted_count = 0
        for sec in sections:
            heading = sec.get("heading", "")
            paragraphs_text = sec.get("paragraphs", [])

            # Tambah heading bab
            if heading:
                new_heading = doc.add_heading(heading, level=2)
                # Pindahkan ke posisi yang tepat
                placeholder_parent.insert(placeholder_index + inserted_count, new_heading._element)
                inserted_count += 1

            # Tambah paragraf-paragraf bab
            for para_text in paragraphs_text:
                new_para = doc.add_paragraph(para_text)
                placeholder_parent.insert(placeholder_index + inserted_count, new_para._element)
                inserted_count += 1
    elif sections:
        # Tag {{SECTIONS}} tidak ada tapi ada data → append di akhir
        warnings.append("Tag {{SECTIONS}} tidak ditemukan. Bab-bab ditambahkan di akhir dokumen.")
        for sec in sections:
            heading = sec.get("heading", "")
            if heading:
                doc.add_heading(heading, level=2)
            for para_text in sec.get("paragraphs", []):
                doc.add_paragraph(para_text)

    doc.save(output_path)


def _replace_placeholder_in_para(para, placeholder: str, replacement: str) -> None:
    """
    Ganti placeholder di dalam paragraf sambil mempertahankan formatting runs.
    Jika placeholder terbagi antar runs, kita gabungkan dulu.
    """
    # Cek apakah placeholder ada di dalam satu run atau terbagi
    full_text = para.text
    if placeholder not in full_text:
        return

    # Strategi: clear semua runs, terapkan teks baru
    # (simple approach: jika hanya teks polos tanpa formatting kompleks)
    new_text = full_text.replace(placeholder, replacement)

    # Simpan formatting dari run pertama
    if para.runs:
        first_run = para.runs[0]
        bold = first_run.bold
        italic = first_run.italic
        font_name = first_run.font.name
        font_size = first_run.font.size

        # Clear semua runs
        for run in para.runs:
            run.text = ""
        
        # Set teks di run pertama
        first_run.text = new_text
        first_run.bold = bold
        first_run.italic = italic
        if font_name:
            first_run.font.name = font_name
        if font_size:
            first_run.font.size = font_size
    else:
        para.add_run(new_text)


def _export_smart_replace(
    structured_doc: dict,
    template_path: str,
    output_path: str,
) -> None:
    """
    Fallback: template tanpa placeholder tag.
    Deteksi paragraf judul/abstrak/kata kunci secara heuristik dan ganti kontennya.
    Strategi ini paling tidak deterministik, tapi lebih baik dari crash.
    """
    doc = Document(template_path)

    title = structured_doc.get("title", "")
    abstract = structured_doc.get("abstract", "")
    keywords = structured_doc.get("keywords", [])
    sections = structured_doc.get("sections", [])
    keywords_str = "; ".join(keywords)

    def replace_para_text(p, new_text: str):
        if p.runs:
            p.runs[0].text = new_text
            for r in list(p.runs)[1:]:
                r.text = ""
        else:
            p.add_run(new_text)

    keyword_para_idx = -1

    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip().lower()

        # Deteksi judul: paragraf pertama yang tidak kosong dan tidak seperti header jurnal
        # (heuristik: panjang > 10 karakter, tidak mengandung "issn", "vol", "no.")
        if (
            title
            and idx < 5
            and len(txt) > 10
            and not any(x in txt for x in ["issn", "vol.", "no.", "jurnal", "journal"])
            and p.text.strip()
        ):
            if not title.lower()[:20] in txt:  # Belum diganti
                replace_para_text(p, title)
                title = ""  # Tandai sudah diganti

        if "kata kunci" in txt or "keywords" in txt:
            keyword_para_idx = idx
            label = "Kata Kunci: " if "kata kunci" in txt else "Keywords: "
            replace_para_text(p, label + keywords_str)

        if "abstrak" in txt or "abstract" in txt:
            if len(p.text.strip()) > 50:
                prefix = "Abstrak— " if "abstrak" in txt else "Abstract— "
                replace_para_text(p, prefix + abstract)
            elif idx + 1 < len(doc.paragraphs):
                next_p = doc.paragraphs[idx + 1]
                if len(next_p.text.strip()) > 50:
                    prefix = "Abstrak— " if "abstrak" in txt else "Abstract— "
                    replace_para_text(next_p, prefix + abstract)

    # Hapus konten setelah kata kunci untuk memberi ruang bab-bab baru
    if keyword_para_idx != -1:
        for p in list(doc.paragraphs)[keyword_para_idx + 1:]:
            p._element.getparent().remove(p._element)

    # Append sections
    for sec in sections:
        heading = sec.get("heading", "")
        if heading:
            doc.add_heading(heading, level=2)
        for para_text in sec.get("paragraphs", []):
            doc.add_paragraph(para_text)

    doc.save(output_path)


def _export_via_jinja(
    structured_doc: dict,
    template_path: str,
    output_path: str,
) -> None:
    """Gunakan docxtpl untuk template Jinja2. Backward compatible dengan format lama."""
    from docxtpl import DocxTemplate, RichText
    
    sections = structured_doc.get("sections", [])
    keywords = structured_doc.get("keywords", [])
    
    def _build_rt(text: str) -> RichText:
        rt = RichText()
        paras = [p.strip() for p in text.split("\n\n") if p.strip()]
        for i, p in enumerate(paras):
            if i > 0:
                rt.add("\n")
            rt.add(p)
        return rt

    doc = DocxTemplate(template_path)
    context = {
        "judul_artikel": structured_doc.get("title", ""),
        "abstrak": _build_rt(structured_doc.get("abstract", "")),
        "kata_kunci": "; ".join(keywords),
        "nama_penulis": structured_doc.get("author", "[Nama Penulis]"),
        "afiliasi": structured_doc.get("affiliation", "[Afiliasi]"),
        "email_korespondensi": structured_doc.get("email", "[Email]"),
        "daftar_bab": [
            {
                "judul_bab": sec.get("heading", ""),
                "isi_bab": _build_rt("\n\n".join(sec.get("paragraphs", []))),
            }
            for sec in sections
        ],
        "daftar_referensi": structured_doc.get("references_formatted", []),
    }
    doc.render(context)
    doc.save(output_path)


def _export_plain_docx(structured_doc: dict, output_path: str) -> None:
    """Buat dokumen .docx standar tanpa template (fallback terakhir)."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = structured_doc.get("title", "")
    abstract = structured_doc.get("abstract", "")
    keywords = structured_doc.get("keywords", [])
    sections = structured_doc.get("sections", [])

    if title:
        h = doc.add_heading(title, level=1)
        h.alignment = 1  # CENTER

    if abstract:
        doc.add_heading("Abstrak", level=2)
        doc.add_paragraph(abstract)

    if keywords:
        p = doc.add_paragraph()
        run = p.add_run("Kata Kunci: ")
        run.bold = True
        p.add_run("; ".join(keywords))

    for sec in sections:
        heading = sec.get("heading", "")
        if heading:
            doc.add_heading(heading, level=2)
        for para_text in sec.get("paragraphs", []):
            doc.add_paragraph(para_text)

    doc.save(output_path)


def structured_doc_from_pipeline(
    title: str,
    abstract: str,
    keywords: list[str],
    sections: list[dict],
    references_formatted: Optional[list[dict]] = None,
    author: str = "[Nama Penulis]",
    affiliation: str = "[Afiliasi]",
    email: str = "[Email Korespondensi]",
) -> dict:
    """
    Helper: Buat structured_doc dict dari output pipeline AI.
    Sehingga pipeline lama tetap compatible dengan injector baru ini.
    """
    return {
        "title": title,
        "abstract": abstract,
        "keywords": keywords,
        "author": author,
        "affiliation": affiliation,
        "email": email,
        "sections": [
            {
                "heading": sec.get("title", sec.get("heading", "")),
                "paragraphs": (
                    [sec["content"]]
                    if isinstance(sec.get("content"), str)
                    else sec.get("paragraphs", [])
                ),
            }
            for sec in sections
        ],
        "references_formatted": references_formatted or [],
    }
