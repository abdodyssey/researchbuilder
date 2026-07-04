"""
DOCX Exporter — Export Artikel ke Format Word (.docx)
======================================================
Mengkonversi hasil pipeline (structured article data atau markdown)
menjadi file .docx yang siap download.

Dua mode export:
1. Template-based (DocxTemplate + Jinja2):
   - Gunakan template .docx user yang sudah punya tag {{ judul_artikel }}, {{ abstrak }}, dll
   - DocxTemplate render otomatis → output rapi sesuai format jurnal user

2. Fallback (python-docx langsung):
   - Jika template tidak punya tag Jinja → parse markdown → inject per-paragraf
   - Tetap gunakan template sebagai "shell" (header, footer, font, margin)
   - Konten di-append setelah section kata kunci

Helper functions:
- build_bab_richtext(): Konversi paragraf markdown → RichText docxtpl (bold/italic)
- extract_template_text(): Ambil teks dari .docx untuk analisis template
- export_to_docx(): Entry point utama, pilih mode template/fallback
- export_markdown_to_docx_fallback(): Fallback mode (parse .md → .docx)
- _add_formatted_text(): Helper untuk render bold/italic di python-docx
"""

import re
from pathlib import Path
import docx
from docxtpl import DocxTemplate, RichText


def build_bab_richtext(isi_bab: str) -> RichText:
    """
    Konversi isi bab yang berupa string dengan pemisah paragraf (\\n\\n)
    menjadi RichText docxtpl, supaya tiap paragraf AI benar-benar jadi
    paragraf baru di Word (bukan newline mentah yang di-collapse).
    """
    rt = RichText()
    paragraphs = [p.strip() for p in isi_bab.split("\n\n") if p.strip()]
    for i, para in enumerate(paragraphs):
        if i > 0:
            rt.add("\n")  # docxtpl RichText newline -> line break asli di Word
            
        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', para)
        for token in tokens:
            if token.startswith("**") and token.endswith("**"):
                rt.add(token[2:-2], bold=True)
            elif token.startswith("*") and token.endswith("*") and len(token) > 2:
                rt.add(token[1:-1], italic=True)
            else:
                rt.add(token)
    return rt


def extract_template_text(template_path: str) -> str:
    """Ambil semua teks dari file .docx template untuk analisis (cek Jinja tags, dll)."""
    try:
        doc = docx.Document(template_path)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception:
        return ""


def export_to_docx(article_data: dict, template_path: str, output_path: str, md_path: str = None) -> str:
    """
    Export artikel ke .docx.

    Dua jalur:
    1. Jika template punya tag Jinja ({{ ... }}) → render via DocxTemplate
    2. Jika tidak → fallback ke export_markdown_to_docx_fallback (parse .md)

    article_data harus punya struktur:
    {
        "judul_artikel": str,
        "nama_penulis": str,
        "afiliasi": str,
        "email_korespondensi": str,
        "abstrak": str,
        "kata_kunci": str,
        "daftar_bab": [
            {"judul_bab": str, "isi_bab": str},  # isi_bab pisah paragraf pakai \n\n
            ...
        ],
        "daftar_referensi": [
            {"teks_sitasi": str},
            ...
        ],
    }
    """

    # Periksa apakah template memiliki tag jinja
    has_jinja_tags = False
    if template_path:
        txt = extract_template_text(template_path)
        if "{{" in txt or "{%" in txt:
            has_jinja_tags = True
            
    if not has_jinja_tags and md_path:
        # Fallback to manual replacement for custom templates
        export_markdown_to_docx_fallback(md_path, output_path, template_path)
        return output_path

    doc = DocxTemplate(template_path)

    context = {
        "judul_artikel": article_data["judul_artikel"],
        "nama_penulis": article_data["nama_penulis"],
        "afiliasi": article_data["afiliasi"],
        "email_korespondensi": article_data["email_korespondensi"],
        "abstrak": build_bab_richtext(article_data["abstrak"]),
        "kata_kunci": article_data["kata_kunci"],
        "daftar_bab": [
            {
                "judul_bab": bab["judul_bab"],
                "isi_bab": build_bab_richtext(bab["isi_bab"]),
            }
            for bab in article_data["daftar_bab"]
        ],
        "daftar_referensi": article_data["daftar_referensi"],
    }

    doc.render(context)
    doc.save(output_path)
    return output_path


def export_markdown_to_docx_fallback(md_path: str, docx_path: str, template_path: str = None) -> None:
    """
    Fallback export: parse markdown → inject ke .docx.

    Digunakan ketika template .docx user TIDAK punya tag Jinja.
    Flow:
    1. Parse YAML frontmatter (title, keywords)
    2. Jika ada template: load, replace title/abstrak/keywords di header template
    3. Hapus body template lama (setelah keywords)
    4. Parse setiap baris markdown → append ke dokumen (heading, list, paragraf)
    5. Bold/italic formatting dipertahankan via _add_formatted_text()
    """
    md_content = Path(md_path).read_text(encoding="utf-8")
    
    # 1. Parse metadata (title, abstract, keywords) from markdown frontmatter
    title = ""
    keywords_list = []
    
    fm_match = re.match(r"^---\s*\n(.*?)\n---\s*\n", md_content, re.DOTALL)
    if fm_match:
        fm_text = fm_match.group(1)
        t_match = re.search(r"^title:\s*[\"']?(.*?)[\"']?\s*$", fm_text, re.MULTILINE)
        if t_match:
            title = t_match.group(1)
        k_match = re.search(r"^keywords:\s*\[?(.*?)\]?\s*$", fm_text, re.MULTILINE)
        if k_match:
            k_raw = k_match.group(1)
            keywords_list = [k.strip(" '\"") for k in k_raw.split(",")]
            
    # Parse abstract from body
    abstract_match = re.search(r"## Abstrak\s*\n(.*?)(?=\n##|$)", md_content, re.DOTALL)
    abstract_text = abstract_match.group(1).strip() if abstract_match else ""
    
    keywords_str = ", ".join(keywords_list)
    
    doc = None
    keyword_idx = -1
    
    # 2. If a docx template is provided, load it and perform smart inline replacement
    if template_path and Path(template_path).suffix.lower() == ".docx":
        try:
            doc = docx.Document(template_path)
            
            # Helper to replace text of a paragraph while keeping its runs/formatting intact
            def replace_paragraph_text_preserving_format(p, new_text):
                if not p.runs:
                    p.add_run(new_text)
                else:
                    p.runs[0].text = new_text
                    # Remove all extra runs so we don't have overlapping old text
                    for r in list(p.runs)[1:]:
                        p._element.remove(r._element)
            
            # Find title paragraph
            # The title is usually the first non-empty paragraph that is not a header/journal metadata
            title_p = None
            for p in doc.paragraphs:
                cleaned = p.text.strip()
                if cleaned and len(cleaned) > 5:
                    if not any(x in cleaned.lower() for x in ["jurnal", "issn", "volume", "halaman", "vol.", "no."]):
                        title_p = p
                        break
            if title_p and title:
                replace_paragraph_text_preserving_format(title_p, title)
                
            # Find and replace Abstrak / Abstract and Keywords / Kata Kunci
            for idx, p in enumerate(list(doc.paragraphs)):
                txt = p.text.lower()
                if "kata kunci" in txt or "keywords" in txt:
                    keyword_idx = idx
                    if "kata kunci" in txt:
                        replace_paragraph_text_preserving_format(p, f"Kata Kunci: {keywords_str}")
                    else:
                        replace_paragraph_text_preserving_format(p, f"Keywords: {keywords_str}")
                elif "abstrak" in txt or "abstract" in txt:
                    # If the paragraph itself contains the abstract text (longer paragraph)
                    if len(p.text.strip()) > 50:
                        prefix = "Abstrak—" if "abstrak" in txt else "Abstract—"
                        replace_paragraph_text_preserving_format(p, f"{prefix}{abstract_text}")
                    # If it's a short header, check the next paragraph for the body
                    elif idx + 1 < len(doc.paragraphs):
                        next_p = doc.paragraphs[idx + 1]
                        if len(next_p.text.strip()) > 50:
                            prefix = "Abstrak—" if "abstrak" in txt else "Abstract—"
                            replace_paragraph_text_preserving_format(next_p, f"{prefix}{abstract_text}")
                            
            # Delete all paragraphs after the keyword paragraph to clean the template body
            # but keep the section break containing the multi-column formatting
            if keyword_idx != -1:
                for p in list(doc.paragraphs)[keyword_idx + 1:]:
                    p._element.getparent().remove(p._element)
            else:
                # If keywords section not found, fallback to clearing all paragraphs
                for p in list(doc.paragraphs):
                    p._element.getparent().remove(p._element)
        except Exception as e:
            print(f"[Warning] Failed to use docx template: {e}")
            doc = docx.Document()
    else:
        doc = docx.Document()
        
    # 3. If no template or template failed, apply standard styling
    if not template_path or not doc.paragraphs:
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = docx.shared.Pt(11)
        
    # 4. Parse the markdown sections and append to the docx document
    in_yaml = False
    for line in md_content.splitlines():
        trimmed = line.strip()
        
        # Skip YAML frontmatter
        if trimmed == "---":
            in_yaml = not in_yaml
            continue
        if in_yaml:
            continue
            
        # Empty lines
        if not trimmed:
            continue
            
        # Skip Abstrak section since it was already replaced in the template header
        # (Only if we are using a template and successfully replaced the header elements)
        if template_path and keyword_idx != -1:
            # Skip Title heading and Abstrak section in the body list
            if trimmed.startswith("# "):
                continue
            if trimmed.startswith("## Abstrak") or trimmed.startswith("## Abstract"):
                # We skip the heading and the subsequent abstract paragraph
                continue
            # Check if this paragraph is the abstract paragraph that follows the heading
            # (We skip it by checking if it's the exact abstract text)
            if abstract_text and trimmed.startswith(abstract_text[:min(50, len(abstract_text))]):
                continue
                
        # Blockquotes (warnings)
        if trimmed.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = docx.shared.Inches(0.4)
            run = p.add_run(trimmed.lstrip(">").strip())
            run.italic = True
            continue
            
        # Headings
        if trimmed.startswith("# "):
            try:
                doc.add_heading(trimmed[2:], level=1)
            except Exception:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = docx.shared.Pt(12)
                p.paragraph_format.space_after = docx.shared.Pt(6)
                run = p.add_run(trimmed[2:])
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = docx.shared.Pt(16)
            continue
        elif trimmed.startswith("## "):
            try:
                doc.add_heading(trimmed[3:], level=2)
            except Exception:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = docx.shared.Pt(12)
                p.paragraph_format.space_after = docx.shared.Pt(4)
                run = p.add_run(trimmed[3:])
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = docx.shared.Pt(13)
            continue
        elif trimmed.startswith("### "):
            try:
                doc.add_heading(trimmed[4:], level=3)
            except Exception:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = docx.shared.Pt(6)
                p.paragraph_format.space_after = docx.shared.Pt(2)
                run = p.add_run(trimmed[4:])
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = docx.shared.Pt(11)
            continue
            
        # Bullet list items
        if trimmed.startswith("- ") or trimmed.startswith("* "):
            try:
                p = doc.add_paragraph(style='List Bullet')
                _add_formatted_text(p, trimmed[2:])
            except Exception:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = docx.shared.Inches(0.25)
                p.paragraph_format.space_after = docx.shared.Pt(3)
                p.add_run("• ")
                _add_formatted_text(p, trimmed[2:])
            continue
            
        # Standard paragraphs
        p = doc.add_paragraph()
        p.paragraph_format.space_after = docx.shared.Pt(6)
        _add_formatted_text(p, line)
        
    doc.save(docx_path)



def _add_formatted_text(paragraph, text: str):
    """Render teks dengan bold (**text**) dan italic (*text*) ke paragraph Word."""
    # Split by bold and italic syntax: **bold** or *italic*
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)
