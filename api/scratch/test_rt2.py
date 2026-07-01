from docxtpl import DocxTemplate, RichText
import docx

def main():
    doc = docx.Document()
    doc.add_paragraph("{{ my_rt }}")
    doc.add_paragraph("{{r my_rt }}")
    doc.add_paragraph("{%r my_rt %}")
    doc.save("scratch/test_rt2_template.docx")
    
    
    tpl = DocxTemplate("scratch/test_rt2_template.docx")
    rt = RichText()
    rt.add("Line 1")
    rt.add("\n")
    rt.add("Line 2")
    
    # Docxtpl replaces {{r my_rt }} ? Let's see.
    try:
        tpl.render({"my_rt": rt})
        tpl.save("scratch/test_rt2_out.docx")
        print("Render successful.")
    except Exception as e:
        print("Error rendering:", e)
        
if __name__ == "__main__":
    main()
